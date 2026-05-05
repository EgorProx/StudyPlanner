import sys
import os
import logging
import database
import notifications
from schedule_manager import ScheduleEntryDialog, WeekCalculator, ScheduleParserBase, SfuScheduleParser
from grades_manager import GradeDialog, GradeStatistics
from ai_chat import AIChatWidget, APIKeyDialog
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PyQt6.QtWidgets import (QApplication, QMainWindow, QMessageBox, QWidget,
                             QVBoxLayout, QHBoxLayout, QListWidget, QLabel,
                             QLineEdit, QPushButton, QTextEdit, QCalendarWidget,
                             QComboBox, QFileDialog, QInputDialog, QFrame, QDialog, QDialogButtonBox, QFormLayout,
                             QListWidgetItem, QTableWidget, QTableWidgetItem as TableItem,
                             QHeaderView, QGroupBox, QGridLayout, QSplitter, QTabWidget,
                             QDoubleSpinBox, QSpinBox, QRadioButton, QButtonGroup,
                             QScrollArea, QStackedWidget, QTableWidgetItem)
from PyQt6.QtCore import Qt, QDate
from PyQt6.QtGui import QTextCharFormat, QColor, QBrush, QFont, QTextCursor, QTextBlockFormat, QKeySequence, QShortcut
from ui_py.ui_main import Ui_MainWindow

logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log', encoding='utf-8'),
        logging.StreamHandler(sys.stdout)
    ]
)
logger = logging.getLogger(__name__)


class TaskDialog(QDialog):
    def __init__(self, parent=None, subjects=None, task_data=None):
        super().__init__(parent)
        self.setWindowTitle('Задача')
        self.resize(300, 200)
        layout = QVBoxLayout(self)
        form_layout = QFormLayout()
        self.input_title = QLineEdit()
        self.input_date = QLineEdit()
        self.input_date.setPlaceholderText('ГГГГ-ММ-ДД')
        self.combo_subject = QComboBox()
        self.subjects_map = {}
        if subjects:
            for sub in subjects:
                self.combo_subject.addItem(sub[1], sub[0])
                self.subjects_map[sub[1]] = sub[0]
        self.combo_subject.addItem('Без предмета', None)
        self.input_desc = QLineEdit()
        form_layout.addRow('Название:', self.input_title)
        form_layout.addRow('Срок:', self.input_date)
        form_layout.addRow('Предмет:', self.combo_subject)
        form_layout.addRow('Описание:', self.input_desc)
        layout.addLayout(form_layout)
        buttons = QDialogButtonBox(QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)
        if task_data:
            self.input_title.setText(task_data[1])
            self.input_date.setText(task_data[2] if task_data[2] else '')
            self.input_desc.setText(task_data[4] if len(task_data) > 4 else '')
            if task_data[3]:
                index = self.combo_subject.findText(task_data[3])
                if index >= 0:
                    self.combo_subject.setCurrentIndex(index)
            else:
                self.combo_subject.setCurrentText('Без предмета')

    def get_data(self):
        subject_id = self.combo_subject.currentData()
        return {
            'title': self.input_title.text(),
            'due_date': self.input_date.text(),
            'subject_id': subject_id,
            'description': self.input_desc.text()
        }


class SearchDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle('Глобальный поиск')
        self.resize(400, 100)
        self.parent_window = parent
        layout = QVBoxLayout(self)
        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText('Введите текст для поиска...')
        self.search_input.returnPressed.connect(self.perform_search)
        layout.addWidget(self.search_input)
        btn_search = QPushButton('Найти')
        btn_search.clicked.connect(self.perform_search)
        layout.addWidget(btn_search)
        self.search_input.setFocus()

    def perform_search(self):
        query = self.search_input.text().strip()
        if not query:
            return
        results = database.global_search(query)
        self.close()
        if not results:
            QMessageBox.information(self.parent_window, 'Поиск', 'Ничего не найдено')
            return
        results_dialog = SearchResultsDialog(self.parent_window, results)
        results_dialog.exec()


class SearchResultsDialog(QDialog):
    def __init__(self, parent=None, results=None):
        super().__init__(parent)
        self.setWindowTitle('Результаты поиска')
        self.resize(400, 300)
        self.parent_window = parent
        layout = QVBoxLayout(self)
        self.list_widget = QListWidget()
        layout.addWidget(self.list_widget)
        self.results_map = {}
        if results:
            for i, item in enumerate(results):
                try:
                    item_id, name, item_type, extra = item
                    if item_type == 'subject':
                        display_text = f'[Предмет] {name}'
                        self.list_widget.addItem(display_text)
                        self.results_map[self.list_widget.count() - 1] = ('subject', item_id)
                    elif item_type == 'task':
                        display_text = f'[Задание] {name}'
                        if extra:
                            display_text += f' (Предмет: {extra})'
                        self.list_widget.addItem(display_text)
                        self.results_map[self.list_widget.count() - 1] = ('task', item_id)
                except Exception as e:
                    logger.error(f'Error processing item {i}: {e}')
        self.list_widget.itemClicked.connect(self.on_item_clicked)
        self.list_widget.itemDoubleClicked.connect(self.on_item_clicked)
        btn_close = QPushButton('Закрыть')
        btn_close.clicked.connect(self.close)
        layout.addWidget(btn_close)

    def on_item_clicked(self, item):
        row = self.list_widget.row(item)
        if row in self.results_map:
            item_type, item_id = self.results_map[row]
            self.parent_window.navigate_to_item(item_type, item_id)
            self.close()


class WishlistDialog(QDialog):
    def __init__(self, parent=None, item_data=None):
        super().__init__(parent)
        self.setWindowTitle('Редактирование желания' if item_data else 'Новое желание')
        self.resize(400, 350)
        layout = QVBoxLayout(self)
        form = QFormLayout()

        self.input_title = QLineEdit()
        self.input_title.setPlaceholderText('Название желания')
        form.addRow('Название:', self.input_title)

        self.input_desc = QLineEdit()
        self.input_desc.setPlaceholderText('Описание...')
        form.addRow('Описание:', self.input_desc)

        self.combo_category = QComboBox()
        self.combo_category.addItems(['Учеба', 'Личное', 'Книги', 'Навыки', 'Проекты', 'Другое'])
        self.combo_category.setEditable(True)
        form.addRow('Категория:', self.combo_category)

        self.combo_priority = QComboBox()
        self.combo_priority.addItems(['Высокий', 'Средний', 'Низкий'])
        self.combo_priority.setCurrentIndex(1)
        form.addRow('Приоритет:', self.combo_priority)

        layout.addLayout(form)

        btn_layout = QHBoxLayout()
        btn_ok = QPushButton('Сохранить')
        btn_ok.clicked.connect(self.accept)
        btn_cancel = QPushButton('Отмена')
        btn_cancel.clicked.connect(self.reject)
        btn_layout.addWidget(btn_ok)
        btn_layout.addWidget(btn_cancel)
        layout.addLayout(btn_layout)

        if item_data:
            self.load_data(item_data)

    def load_data(self, data):
        item_id, title, description, category, priority, status, created_at = data
        self.input_title.setText(title or '')
        self.input_desc.setText(description or '')
        if category:
            idx = self.combo_category.findText(category)
            if idx >= 0:
                self.combo_category.setCurrentIndex(idx)
            else:
                self.combo_category.setCurrentText(category)
        priority_map = {'high': 'Высокий', 'medium': 'Средний', 'low': 'Низкий'}
        priority_name = priority_map.get(priority, 'Средний')
        idx = self.combo_priority.findText(priority_name)
        if idx >= 0:
            self.combo_priority.setCurrentIndex(idx)

    def get_data(self):
        priority_map = {'Высокий': 'high', 'Средний': 'medium', 'Низкий': 'low'}
        return {
            'title': self.input_title.text(),
            'description': self.input_desc.text(),
            'category': self.combo_category.currentText(),
            'priority': priority_map.get(self.combo_priority.currentText(), 'medium')
        }


STYLES = {
    "light": (
        "QMainWindow, QWidget { background-color: #f5f5f5; color: #333; font-family: 'Segoe UI', sans-serif; }"
        "QListWidget { background-color: #ffffff; border: none; border-right: 1px solid #ddd; font-size: 16px; }"
        "QListWidget::item { padding: 15px; border-bottom: 1px solid #f0f0f0; }"
        "QListWidget::item:selected { background-color: #e3f2fd; color: #000; font-weight: bold; }"
        "QPushButton { background-color: #007bff; color: white; border-radius: 6px; padding: 10px; font-weight: bold; border: none; }"
        "QPushButton:hover { background-color: #0056b3; }"
        "QLineEdit, QTextEdit, QComboBox { background-color: white; border: 1px solid #ccc; border-radius: 6px; padding: 8px; color: #333; }"
        "QCalendarWidget QTableView { background-color: white; color: #333; selection-background-color: #007bff; selection-color: white; }"
        "QCalendarWidget QToolButton { color: #333; background-color: transparent; }"
        "QCalendarWidget QWidget#qt_calendar_navigationbar { background-color: #e0e0e0; }"
        "QTableWidget { background-color: white; gridline-color: #ddd; border: 1px solid #ddd; }"
        "QTableWidget::item:selected { background-color: #e3f2fd; color: #000; }"
        "QHeaderView::section { background-color: #f0f0f0; padding: 8px; font-weight: bold; border: 1px solid #ddd; }"
        "QGroupBox { font-weight: bold; border: 1px solid #ccc; border-radius: 6px; margin-top: 10px; padding-top: 10px; }"
        "QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 5px; }"
    ),
    "dark": (
        "QMainWindow, QWidget { background-color: #121212; color: #e0e0e0; font-family: 'Segoe UI', sans-serif; }"
        "QListWidget { background-color: #1e1e1e; border: none; border-right: 1px solid #333; font-size: 16px; }"
        "QListWidget::item { padding: 15px; border-bottom: 1px solid #2c2c2c; }"
        "QListWidget::item:selected { background-color: #3a7bd5; color: white; font-weight: bold; }"
        "QPushButton { background-color: #3a7bd5; color: white; border-radius: 6px; padding: 10px; font-weight: bold; border: none; }"
        "QPushButton:hover { background-color: #5a95e5; }"
        "QLineEdit, QTextEdit, QComboBox { background-color: #2d2d2d; border: 1px solid #444; border-radius: 6px; padding: 8px; color: white; }"
        "QCalendarWidget { background-color: #121212; }"
        "QCalendarWidget QTableView { background-color: #1e1e1e; color: #e0e0e0; selection-background-color: #3a7bd5; selection-color: white; gridline-color: #444; alternate-background-color: #252525; outline: 0px; }"
        "QCalendarWidget QTableView::item:selected { background-color: #3a7bd5; color: white; }"
        "QCalendarWidget QTableView::item:hover { background-color: #2a2a2a; }"
        "QCalendarWidget QToolButton { color: #e0e0e0; background-color: #252525; border: 1px solid #444; border-radius: 4px; margin: 2px; padding: 4px; font-weight: bold; }"
        "QCalendarWidget QToolButton:hover { background-color: #3a3a3a; }"
        "QCalendarWidget QToolButton::menu-indicator { image: none; }"
        "QCalendarWidget QWidget#qt_calendar_navigationbar { background-color: #252525; }"
        "QCalendarWidget QSpinBox { background-color: #2d2d2d; color: white; selection-color: white; selection-background-color: #3a7bd5; border: 1px solid #444; padding: 2px; border-radius: 4px; }"
        "QCalendarWidget QSpinBox::up-button, QCalendarWidget QSpinBox::down-button { background-color: #3a3a3a; border: none; width: 15px; }"
        "QCalendarWidget QSpinBox::up-button:hover, QCalendarWidget QSpinBox::down-button:hover { background-color: #4a4a4a; }"
        "QTableWidget { background-color: #1e1e1e; gridline-color: #444; border: 1px solid #444; color: #e0e0e0; }"
        "QTableWidget::item:selected { background-color: #3a7bd5; color: white; }"
        "QHeaderView::section { background-color: #2d2d2d; padding: 8px; font-weight: bold; border: 1px solid #444; color: #e0e0e0; }"
        "QGroupBox { font-weight: bold; border: 1px solid #444; border-radius: 6px; margin-top: 10px; padding-top: 10px; color: #e0e0e0; }"
        "QGroupBox::title { subcontrol-origin: margin; left: 10px; padding: 0 5px; }"
    )
}


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        logger.info('=== START ===')
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        database.init_db()
        self.current_theme = database.get_setting('theme', 'light')
        self.apply_theme(self.current_theme)
        self.storage_path = database.get_setting('path', '')
        self.current_note_path = None
        self.is_subj_reverse = False
        self.is_task_reverse = False
        self.setup_subjects_ui()
        self.setup_schedule_ui()
        self.setup_settings_ui()
        self.setup_calendar_ui()
        self.setup_tasks_ui()
        self.setup_notes_ui()
        self.setup_grades_ui()
        self.setup_ai_ui()
        self.setup_wishlist_ui()
        self.setup_menu_logic()
        self.setup_global_search()
        self.load_subjects()
        self.load_tasks()
        self.update_calendar_deadlines()
        logger.info('=== INIT DONE ===')
        self.notification_manager = notifications.NotificationManager(self)
        self.notification_manager.initialize()
        self.setup_notifications_menu()

    def apply_theme(self, theme_name):
        app = QApplication.instance()
        app.setStyleSheet(STYLES.get(theme_name, STYLES['light']))

    def setup_menu_logic(self):
        self.ui.menuList.addItem('Предметы')
        self.ui.menuList.addItem('Расписание')
        self.ui.menuList.addItem('Календарь')
        self.ui.menuList.addItem('Настройки')
        self.ui.menuList.addItem('Задания')
        self.ui.menuList.addItem('Блокнот')
        self.ui.menuList.addItem('Оценки')
        self.ui.menuList.addItem('ИИ-ассистент')
        self.ui.menuList.addItem('Желания')
        self.ui.menuList.currentRowChanged.connect(self.change_page)

    def setup_global_search(self):
        search_menu = self.ui.menubar.addMenu('Поиск')
        search_action = search_menu.addAction('Глобальный поиск')
        search_action.triggered.connect(self.open_search_dialog)
        self.search_shortcut = QShortcut(QKeySequence('Ctrl+F'), self)
        self.search_shortcut.activated.connect(self.open_search_dialog)

    def open_search_dialog(self):
        dialog = SearchDialog(self)
        dialog.exec()

    def navigate_to_item(self, item_type, item_id):
        logger.debug(f'navigate_to_item: type={item_type}, id={item_id}')
        if item_type == 'subject':
            self.ui.menuList.setCurrentRow(0)
            self.ui.pagesStack.setCurrentWidget(self.ui.page_subjects)
            self.load_subjects()
            for i in range(self.subjects_list.count()):
                item = self.subjects_list.item(i)
                if item.data(Qt.ItemDataRole.UserRole) == item_id:
                    self.subjects_list.setCurrentItem(item)
                    self.show_subject_details(item)
                    break
        elif item_type == 'task':
            self.ui.menuList.setCurrentRow(4)
            self.ui.pagesStack.setCurrentWidget(self.ui.page_tasks)
            self.combo_task_status.setCurrentIndex(self.combo_task_status.findData('all'))
            self.load_tasks()
            for i in range(self.tasks_list.count()):
                item = self.tasks_list.item(i)
                if item.data(Qt.ItemDataRole.UserRole) == item_id:
                    self.tasks_list.setCurrentItem(item)
                    self.show_task_details(item)
                    break

    def get_task_status_info(self, due_date_str, completed, status):
        if completed:
            return 'Выполнено', QColor(0, 128, 0)
        if status == 'archived':
            return 'В архиве', QColor(128, 128, 128)
        if not due_date_str:
            return 'Без даты', QColor(128, 128, 128)
        try:
            due_date = QDate.fromString(due_date_str, 'yyyy-MM-dd')
            today = QDate.currentDate()
            if due_date < today:
                return 'Просрочено', QColor(220, 20, 60)
            elif due_date <= today.addDays(3):
                return 'Скоро дедлайн', QColor(255, 140, 0)
            else:
                default_color = QColor(0, 0, 0) if self.current_theme == 'light' else QColor(255, 255, 255)
                return 'Активно', default_color
        except:
            return 'Неизвестно', QColor(128, 128, 128)

    def change_page(self, index):
        logger.debug(f'change_page: index={index}')
        if index == 0:
            self.ui.pagesStack.setCurrentWidget(self.ui.page_subjects)
            self.load_subjects()
        elif index == 1:
            self.ui.pagesStack.setCurrentWidget(self.ui.page_schedule)
            self.load_schedule()
        elif index == 2:
            self.ui.pagesStack.setCurrentWidget(self.ui.page_calendar)
            self.update_calendar_deadlines()
        elif index == 3:
            self.ui.pagesStack.setCurrentWidget(self.ui.page_settings)
            self.load_settings_data()
        elif index == 4:
            self.ui.pagesStack.setCurrentWidget(self.ui.page_tasks)
            self.load_tasks()
        elif index == 5:
            self.ui.pagesStack.setCurrentWidget(self.ui.page_notes)
            self.update_note_path_label()
        elif index == 6:
            self.ui.pagesStack.setCurrentWidget(self.ui.page_grades)
            self.load_grades()
        elif index == 7:
            self.ui.pagesStack.setCurrentWidget(self.ui.page_ai)
            self.ai_chat.load_settings()
        elif index == 8:
            self.ui.pagesStack.setCurrentWidget(self.ui.page_wishlist)
            self.load_wishlist()

    def setup_subjects_ui(self):
        logger.debug('setup_subjects_ui')
        page = self.ui.page_subjects
        layout = QHBoxLayout(page)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)
        left_layout = QVBoxLayout()
        left_layout.addWidget(QLabel('<h2>Мои предметы</h2>'))
        sort_layout = QHBoxLayout()
        sort_layout.addWidget(QLabel('Сортировка:'))
        self.combo_subj_sort = QComboBox()
        self.combo_subj_sort.addItem('По названию', 'name')
        self.combo_subj_sort.addItem('По преподавателю', 'teacher')
        self.combo_subj_sort.addItem('По кабинету', 'room')
        self.combo_subj_sort.currentTextChanged.connect(self.load_subjects)
        sort_layout.addWidget(self.combo_subj_sort)
        self.btn_subj_reverse = QPushButton('⇅')
        self.btn_subj_reverse.setFixedWidth(30)
        self.btn_subj_reverse.setToolTip('Изменить порядок')
        self.btn_subj_reverse.clicked.connect(self.toggle_subject_sort)
        sort_layout.addWidget(self.btn_subj_reverse)
        sort_layout.addStretch()
        left_layout.addLayout(sort_layout)
        self.subjects_list = QListWidget()
        self.subjects_list.itemClicked.connect(self.show_subject_details)
        left_layout.addWidget(self.subjects_list)
        btn_layout = QHBoxLayout()
        self.btn_add = QPushButton('Добавить')
        self.btn_edit = QPushButton('Изменить')
        self.btn_del = QPushButton('Удалить')
        self.btn_add.clicked.connect(self.add_subject)
        self.btn_edit.clicked.connect(self.edit_subject)
        self.btn_del.clicked.connect(self.delete_subject)
        btn_layout.addWidget(self.btn_add)
        btn_layout.addWidget(self.btn_edit)
        btn_layout.addWidget(self.btn_del)
        self.btn_clear_subjects = QPushButton('Очистить всё')
        self.btn_clear_subjects.setToolTip('Удалить все предметы, расписание и оценки')
        self.btn_clear_subjects.setStyleSheet('QPushButton { color: #dc3545; } QPushButton:hover { background-color: #dc3545; color: white; }')
        self.btn_clear_subjects.clicked.connect(self.clear_all_subjects)
        btn_layout.addWidget(self.btn_clear_subjects)
        left_layout.addLayout(btn_layout)
        right_layout = QVBoxLayout()
        right_layout.addWidget(QLabel('<h2>Детали</h2>'))
        info_frame = QFrame()
        info_frame.setFrameShape(QFrame.Shape.StyledPanel)
        info_layout = QVBoxLayout(info_frame)
        self.lbl_name = QLabel('Название: -')
        self.lbl_teacher = QLabel('Преподаватель: -')
        self.lbl_room = QLabel('Кабинет: -')
        info_layout.addWidget(self.lbl_name)
        info_layout.addWidget(self.lbl_teacher)
        info_layout.addWidget(self.lbl_room)
        info_layout.addSpacing(10)
        info_layout.addWidget(QLabel('<b>Описание:</b>'))
        self.txt_desc = QLabel('-')
        self.txt_desc.setWordWrap(True)
        self.txt_desc.setAlignment(Qt.AlignmentFlag.AlignTop)
        info_layout.addWidget(self.txt_desc)
        right_layout.addWidget(info_frame)
        right_layout.addStretch()
        layout.addLayout(left_layout, 30)
        layout.addLayout(right_layout, 70)

    def toggle_subject_sort(self):
        logger.debug('toggle_subject_sort')
        self.is_subj_reverse = not self.is_subj_reverse
        self.load_subjects()

    def load_subjects(self):
        logger.debug('load_subjects')
        try:
            self.subjects_list.clear()
            sort_mode = self.combo_subj_sort.currentData()
            if not sort_mode:
                sort_mode = 'name'
            data = database.get_all_subjects(sort_by=sort_mode, reverse=self.is_subj_reverse)
            for row in data:
                if len(row) < 5:
                    continue
                display_text = f'{row[1]} (каб. {row[3] or "-"})'
                item = QListWidgetItem(display_text)
                item.setData(Qt.ItemDataRole.UserRole, row[0])
                self.subjects_list.addItem(item)
            logger.info(f'load_subjects: {self.subjects_list.count()} subjects')
        except Exception as e:
            logger.error(f'load_subjects: {e}', exc_info=True)

    def show_subject_details(self, item):
        logger.debug(f'show_subject_details: item={item}')
        try:
            if not item:
                return
            sid = item.data(Qt.ItemDataRole.UserRole)
            if not sid:
                return
            data = database.get_subject_details(sid)
            if not data or len(data) < 5:
                return
            self.lbl_name.setText(f'<b>Название:</b> {data[1] or "-"}')
            self.lbl_teacher.setText(f'<b>Преподаватель:</b> {data[2] or "-"}')
            self.lbl_room.setText(f'<b>Кабинет:</b> {data[3] or "-"}')
            self.txt_desc.setText(data[4] if data[4] else 'Нет описания')
        except Exception as e:
            logger.error(f'show_subject_details: {e}', exc_info=True)

    def add_subject(self):
        logger.debug('add_subject')
        try:
            name, ok = QInputDialog.getText(self, 'Новый предмет', 'Название:')
            if ok and name:
                teacher, ok1 = QInputDialog.getText(self, 'Новый предмет', 'Преподаватель:')
                room, ok2 = QInputDialog.getText(self, 'Новый предмет', 'Кабинет:')
                desc, ok3 = QInputDialog.getText(self, 'Новый предмет', 'Описание:', text='')
                database.add_subject(name, teacher or '', room or '', desc or '')
                self.load_subjects()
                logger.info(f'add_subject: {name} added')
            elif ok and not name:
                QMessageBox.warning(self, 'Ошибка', 'Название предмета не может быть пустым')
        except Exception as e:
            logger.error(f'add_subject: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', f'Произошла ошибка: {str(e)}')

    def edit_subject(self):
        logger.debug('edit_subject')
        try:
            item = self.subjects_list.currentItem()
            if not item:
                return
            sid = item.data(Qt.ItemDataRole.UserRole)
            data = database.get_subject_details(sid)
            if not data or len(data) < 5:
                QMessageBox.warning(self, 'Ошибка', 'Предмет не найден')
                return
            name, ok = QInputDialog.getText(self, 'Редактирование', 'Название:', text=data[1] or '')
            if ok:
                teacher, ok1 = QInputDialog.getText(self, 'Редактирование', 'Преподаватель:', text=data[2] or '')
                room, ok2 = QInputDialog.getText(self, 'Редактирование', 'Кабинет:', text=data[3] or '')
                desc, ok3 = QInputDialog.getText(self, 'Редактирование', 'Описание:', text=data[4] or '')
                database.update_subject(sid, name or '', teacher or '', room or '', desc or '')
                self.load_subjects()
                for i in range(self.subjects_list.count()):
                    new_item = self.subjects_list.item(i)
                    if new_item.data(Qt.ItemDataRole.UserRole) == sid:
                        self.subjects_list.setCurrentItem(new_item)
                        self.show_subject_details(new_item)
                        break
        except Exception as e:
            logger.error(f'edit_subject: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', f'Произошла ошибка: {str(e)}')

    def delete_subject(self):
        logger.debug('delete_subject')
        try:
            item = self.subjects_list.currentItem()
            if not item:
                return
            sid = item.data(Qt.ItemDataRole.UserRole)
            reply = QMessageBox.question(self, 'Удаление', 'Удалить предмет? Задачи останутся без привязки.',
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.Yes:
                database.delete_subject(sid)
                self.load_subjects()
                self.lbl_name.setText('Название: -')
                self.lbl_teacher.setText('Преподаватель: -')
                self.lbl_room.setText('Кабинет: -')
                self.txt_desc.setText('-')
        except Exception as e:
            logger.error(f'delete_subject: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', f'Произошла ошибка: {str(e)}')

    def clear_all_subjects(self):
        logger.debug('clear_all_subjects')
        try:
            reply = QMessageBox.warning(
                self,
                'Очистка предметов',
                'Удалить ВСЕ предметы, расписание и оценки?\n\nЭто действие невозможно отменить!',
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
                QMessageBox.StandardButton.No
            )
            if reply != QMessageBox.StandardButton.Yes:
                return
            database.clear_all_subjects()
            self.load_subjects()
            self.lbl_name.setText('Название: -')
            self.lbl_teacher.setText('Преподаватель: -')
            self.lbl_room.setText('Кабинет: -')
            self.txt_desc.setText('-')
            QMessageBox.information(self, 'Готово', 'Все предметы удалены')
        except Exception as e:
            logger.error(f'clear_all_subjects: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', str(e))

    def setup_schedule_ui(self):
        page = self.ui.page_schedule
        layout = QVBoxLayout(page)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)

        top_layout = QHBoxLayout()
        top_layout.addWidget(QLabel('<h2>Расписание занятий</h2>'))

        week_group = QGroupBox('')
        week_layout = QHBoxLayout(week_group)
        week_layout.setContentsMargins(8, 4, 8, 4)
        self.lbl_current_week = QLabel()
        self.lbl_current_week.setMinimumWidth(80)
        self.update_week_label()
        week_layout.addWidget(self.lbl_current_week)
        week_group.setFixedWidth(100)
        top_layout.addWidget(week_group)

        top_layout.addStretch()

        self.combo_week_filter = QComboBox()
        self.combo_week_filter.addItem('Все недели', 'all')
        self.combo_week_filter.addItem('Четная', 'even')
        self.combo_week_filter.addItem('Нечетная', 'odd')
        self.combo_week_filter.currentIndexChanged.connect(self.load_schedule)
        top_layout.addWidget(QLabel('Фильтр:'))
        top_layout.addWidget(self.combo_week_filter)
        layout.addLayout(top_layout)

        self.schedule_table = QTableWidget()
        self.schedule_table.setColumnCount(8)
        headers = ['Время', 'Пн', 'Вт', 'Ср', 'Чт', 'Пт', 'Сб', 'Вс']
        self.schedule_table.setHorizontalHeaderLabels(headers)
        self.schedule_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.schedule_table.verticalHeader().setVisible(False)
        self.schedule_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectItems)
        self.schedule_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.schedule_table.cellClicked.connect(self.on_schedule_cell_clicked)
        layout.addWidget(self.schedule_table)

        btn_layout = QHBoxLayout()
        self.btn_add_schedule = QPushButton('Добавить занятие')
        self.btn_edit_schedule = QPushButton('Изменить')
        self.btn_del_schedule = QPushButton('Удалить')
        self.btn_import_schedule = QPushButton('Импорт из СФУ')
        self.btn_add_schedule.clicked.connect(self.add_schedule_entry)
        self.btn_edit_schedule.clicked.connect(self.edit_schedule_entry)
        self.btn_del_schedule.clicked.connect(self.delete_schedule_entry)
        self.btn_import_schedule.clicked.connect(self.import_schedule_from_sfu)
        btn_layout.addWidget(self.btn_add_schedule)
        btn_layout.addWidget(self.btn_edit_schedule)
        btn_layout.addWidget(self.btn_del_schedule)
        btn_layout.addWidget(self.btn_import_schedule)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        self.lbl_schedule_details = QLabel('Выберите ячейку для просмотра деталей')
        self.lbl_schedule_details.setStyleSheet('padding: 10px; border: 1px solid #ccc; border-radius: 4px;')
        layout.addWidget(self.lbl_schedule_details)

    def update_week_label(self):
        week_type = WeekCalculator.get_current_week_parity()
        week_name = 'Четная' if week_type == 'even' else 'Нечетная'
        self.lbl_current_week.setText(f'<b>{week_name}</b>')

    def load_schedule(self):
        logger.debug('load_schedule')
        try:
            week_filter = self.combo_week_filter.currentData()
            if week_filter == 'all':
                week_filter = None
            data = database.get_schedule(week_type=week_filter)
            self.schedule_map = {}
            for row in data:
                entry_id = row[0]
                day = row[2]
                start_time = row[4]
                subject_name = row[9]
                room = row[6] or ''
                lesson_type = row[7] or ''
                week_type = row[3]
                key = (day, start_time)
                if key not in self.schedule_map:
                    self.schedule_map[key] = []
                self.schedule_map[key].append({
                    'id': entry_id,
                    'subject': subject_name,
                    'room': room,
                    'type': lesson_type,
                    'week': week_type,
                    'end_time': row[5] or ''
                })
            self.refresh_schedule_table()
        except Exception as e:
            logger.error(f'load_schedule: {e}', exc_info=True)

    def refresh_schedule_table(self):
        self.schedule_table.setRowCount(0)
        # ИСПРАВЛЕНО: берем времена из ключей schedule_map, а не из записей
        all_times = set()
        for (day, start_time) in self.schedule_map.keys():
            all_times.add(start_time)

        if not all_times:
            time_slots = ['08:00', '09:00', '10:00', '11:00', '12:00',
                          '13:00', '14:00', '15:00', '16:00', '17:00',
                          '18:00', '19:00', '20:00']
        else:
            time_slots = sorted(all_times)

        self.schedule_table.setRowCount(len(time_slots))
        for row_idx, time_str in enumerate(time_slots):
            time_item = QTableWidgetItem(time_str)
            time_item.setFlags(time_item.flags() & ~Qt.ItemFlag.ItemIsEditable)
            self.schedule_table.setItem(row_idx, 0, time_item)
            for day in range(7):
                key = (day, time_str)
                if key in self.schedule_map:
                    entries = self.schedule_map[key]
                    texts = []
                    for entry in entries:
                        week_marker = ''
                        if entry['week'] == 'even':
                            week_marker = ' (чет)'
                        elif entry['week'] == 'odd':
                            week_marker = ' (нечет)'
                        text = f"{entry['subject']}{week_marker}\n{entry['type']}\n{entry['room']}"
                        texts.append(text)
                    cell_text = '\n---\n'.join(texts)
                    item = QTableWidgetItem(cell_text)
                    item.setData(Qt.ItemDataRole.UserRole, [e['id'] for e in entries])
                    if len(entries) == 1:
                        if entry['week'] == 'even':
                            item.setBackground(QBrush(QColor(200, 230, 255)))
                        elif entry['week'] == 'odd':
                            item.setBackground(QBrush(QColor(255, 230, 200)))
                        else:
                            item.setBackground(QBrush(QColor(220, 255, 220)))
                    self.schedule_table.setItem(row_idx, day + 1, item)
        self.schedule_table.resizeRowsToContents()

    def on_schedule_cell_clicked(self, row, column):
        if column == 0:
            return
        item = self.schedule_table.item(row, column)
        if not item:
            self.lbl_schedule_details.setText('Нет занятий в это время')
            return
        entry_ids = item.data(Qt.ItemDataRole.UserRole)
        if not entry_ids:
            return
        details_text = ''
        for entry_id in entry_ids:
            data = database.get_schedule_by_id(entry_id)
            if data:
                week_name = WeekCalculator.format_week_type(data[3])
                details_text += f"<b>{data[9]}</b><br>"
                details_text += f"Тип: {data[7] or '-'}<br>"
                details_text += f"Время: {data[4] or '-'} - {data[5] or '-'}<br>"
                details_text += f"Аудитория: {data[6] or '-'}<br>"
                details_text += f"Неделя: {week_name}<br>"
                details_text += f"Преподаватель: {data[10] or '-'}<br>"
                details_text += '<hr>'
        self.lbl_schedule_details.setText(details_text)

    def add_schedule_entry(self):
        logger.debug('add_schedule_entry')
        try:
            subjects = database.get_all_subjects()
            if not subjects:
                QMessageBox.warning(self, 'Ошибка', 'Сначала добавьте предметы в разделе Предметы')
                return
            dlg = ScheduleEntryDialog(self, subjects)
            if dlg.exec():
                data = dlg.get_data()
                database.add_schedule_entry(**data)
                self.load_schedule()
        except Exception as e:
            logger.error(f'add_schedule_entry: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', f'Произошла ошибка: {str(e)}')

    def edit_schedule_entry(self):
        logger.debug('edit_schedule_entry')
        try:
            current = self.schedule_table.currentItem()
            if not current:
                QMessageBox.warning(self, 'Ошибка', 'Выберите занятие в таблице')
                return
            entry_ids = current.data(Qt.ItemDataRole.UserRole)
            if not entry_ids or len(entry_ids) == 0:
                return
            entry_id = entry_ids[0]
            data = database.get_schedule_by_id(entry_id)
            if not data:
                return
            subjects = database.get_all_subjects()
            dlg = ScheduleEntryDialog(self, subjects, data)
            if dlg.exec():
                new_data = dlg.get_data()
                database.update_schedule_entry(entry_id, **new_data)
                self.load_schedule()
        except Exception as e:
            logger.error(f'edit_schedule_entry: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', f'Произошла ошибка: {str(e)}')

    def delete_schedule_entry(self):
        logger.debug('delete_schedule_entry')
        try:
            current = self.schedule_table.currentItem()
            if not current:
                QMessageBox.warning(self, 'Ошибка', 'Выберите занятие в таблице')
                return
            entry_ids = current.data(Qt.ItemDataRole.UserRole)
            if not entry_ids or len(entry_ids) == 0:
                return
            entry_id = entry_ids[0]
            reply = QMessageBox.question(self, 'Удаление', 'Удалить занятие из расписания?',
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.Yes:
                database.delete_schedule_entry(entry_id)
                self.load_schedule()
                self.lbl_schedule_details.setText('Выберите ячейку для просмотра деталей')
        except Exception as e:
            logger.error(f'delete_schedule_entry: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', f'Произошла ошибка: {str(e)}')

    def import_schedule_from_sfu(self):
        logger.debug('import_schedule_from_sfu')
        try:
            source_type, ok = QInputDialog.getItem(
                self,
                'Импорт расписания СФУ',
                'Источник:',
                ['HTML-файл страницы расписания', 'URL страницы расписания'],
                0,
                False
            )
            if not ok:
                return

            parser = SfuScheduleParser()
            if source_type.startswith('HTML'):
                filepath, _ = QFileDialog.getOpenFileName(
                    self,
                    'Выберите HTML-файл расписания',
                    '',
                    'HTML и текстовые файлы (*.html *.htm *.txt);;Все файлы (*)'
                )
                if not filepath:
                    return
                entries = parser.parse_file(filepath)
            else:
                url, ok = QInputDialog.getText(self, 'URL расписания СФУ', 'Ссылка на страницу:')
                if not ok or not url.strip():
                    return
                entries = parser.parse_url(url.strip())

            if not entries:
                QMessageBox.warning(self, 'Импорт расписания', 'В выбранном источнике не найдено занятий.')
                return

            reply = QMessageBox.question(
                self,
                'Импорт расписания',
                f'Найдено занятий: {len(entries)}.\nОчистить текущее расписание перед импортом?',
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.No:
                pass
            if reply == QMessageBox.StandardButton.Yes:
                database.clear_schedule()

            imported, errors = parser.import_to_database(entries, create_missing_subjects=True)
            self.load_subjects()
            self.load_schedule()

            message = f'Импортировано занятий: {imported} из {len(entries)}.'
            if errors:
                message += '\n\nОшибки:\n' + '\n'.join(errors[:5])
                if len(errors) > 5:
                    message += f'\n...и еще {len(errors) - 5}'
            QMessageBox.information(self, 'Импорт расписания', message)
        except Exception as e:
            logger.error(f'import_schedule_from_sfu: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', f'Не удалось импортировать расписание: {str(e)}')

    def setup_ai_ui(self):
        self.ai_chat = AIChatWidget(self)
        page = self.ui.page_ai
        layout = QVBoxLayout(page)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.addWidget(self.ai_chat)

    def setup_settings_ui(self):
        page = self.ui.page_settings
        layout = QVBoxLayout(page)
        layout.setContentsMargins(40, 40, 40, 40)
        layout.setSpacing(20)
        layout.addWidget(QLabel('<h2>Настройки</h2>'))
        layout.addWidget(QLabel('Оформление:'))
        self.combo_theme = QComboBox()
        self.combo_theme.addItems(['light', 'dark'])
        self.combo_theme.currentTextChanged.connect(self.change_theme)
        layout.addWidget(self.combo_theme)
        layout.addSpacing(20)
        layout.addWidget(QLabel('Рабочая папка:'))
        path_layout = QHBoxLayout()
        self.edit_path = QLineEdit()
        self.edit_path.setReadOnly(True)
        btn_browse = QPushButton('Обзор...')
        btn_browse.clicked.connect(self.browse_folder)
        path_layout.addWidget(self.edit_path)
        path_layout.addWidget(btn_browse)
        layout.addLayout(path_layout)
        layout.addSpacing(20)
        layout.addWidget(QLabel('OpenRouter API:'))
        api_layout = QHBoxLayout()
        self.edit_api_key = QLineEdit()
        self.edit_api_key.setPlaceholderText('Введите API-ключ OpenRouter...')
        self.edit_api_key.setEchoMode(QLineEdit.EchoMode.Password)
        self.edit_api_key.textChanged.connect(self.on_api_key_changed)
        api_layout.addWidget(self.edit_api_key)
        self.btn_api_key = QPushButton('Показать/скрыть')
        self.btn_api_key.clicked.connect(self.toggle_api_key_visibility)
        api_layout.addWidget(self.btn_api_key)
        layout.addLayout(api_layout)
        layout.addStretch()

    def load_settings_data(self):
        self.combo_theme.setCurrentText(self.current_theme)
        self.edit_path.setText(self.storage_path)
        self.edit_api_key.setText(database.get_setting('openrouter_api_key', ''))

    def on_api_key_changed(self, text):
        database.save_setting('openrouter_api_key', text.strip())

    def toggle_api_key_visibility(self):
        if self.edit_api_key.echoMode() == QLineEdit.EchoMode.Password:
            self.edit_api_key.setEchoMode(QLineEdit.EchoMode.Normal)
        else:
            self.edit_api_key.setEchoMode(QLineEdit.EchoMode.Password)

    def change_theme(self, theme):
        logger.debug(f'change_theme: {theme}')
        self.current_theme = theme
        self.apply_theme(theme)
        database.save_setting('theme', theme)
        self.load_tasks()

    def browse_folder(self):
        folder = QFileDialog.getExistingDirectory(self, 'Выберите папку')
        if folder:
            logger.info(f'browse_folder: {folder}')
            self.storage_path = folder
            self.edit_path.setText(folder)
            database.save_setting('path', folder)

    def setup_calendar_ui(self):
        page = self.ui.page_calendar
        layout = QVBoxLayout(page)
        layout.setContentsMargins(40, 40, 40, 40)
        layout.addWidget(QLabel('<h2>Календарь дедлайнов</h2>'))
        self.calendar = QCalendarWidget()
        self.calendar.setGridVisible(True)
        layout.addWidget(self.calendar)
        self.lbl_date = QLabel('Выберите дату')
        self.lbl_date.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.lbl_date)
        self.calendar.clicked.connect(self.on_date_click)

    def update_calendar_deadlines(self):
        logger.debug('update_calendar_deadlines')
        try:
            self.calendar.setDateTextFormat(QDate(), QTextCharFormat())
            tasks = database.get_all_tasks()
            for task in tasks:
                date_str = task[2]
                if date_str:
                    try:
                        y, m, d = map(int, date_str.split('-'))
                        qdate = QDate(y, m, d)
                        fmt = QTextCharFormat()
                        fmt.setBackground(QColor('#ffcccc'))
                        if self.current_theme == 'dark':
                            fmt.setForeground(QBrush(QColor('white')))
                        self.calendar.setDateTextFormat(qdate, fmt)
                    except ValueError:
                        pass
        except Exception as e:
            logger.error(f'update_calendar_deadlines: {e}', exc_info=True)

    def on_date_click(self, date):
        logger.debug(f'on_date_click: {date.toString("yyyy-MM-dd")}')
        try:
            date_str = date.toString('yyyy-MM-dd')
            tasks = database.get_all_tasks()
            task_names = [t[1] for t in tasks if t[2] == date_str]
            if task_names:
                self.lbl_date.setText(f'<b>{date.toString("dddd, d MMMM yyyy")}</b><br>Задачи: {", ".join(task_names)}')
            else:
                self.lbl_date.setText(f'{date.toString("dddd, d MMMM yyyy")}<br>Нет задач')
        except Exception as e:
            logger.error(f'on_date_click: {e}', exc_info=True)

    def setup_tasks_ui(self):
        page = self.ui.page_tasks
        layout = QHBoxLayout(page)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)
        left_layout = QVBoxLayout()
        left_layout.addWidget(QLabel('<h2>Список заданий</h2>'))
        status_layout = QHBoxLayout()
        status_layout.addWidget(QLabel('Показать:'))
        self.combo_task_status = QComboBox()
        self.combo_task_status.addItem('Активные', 'active')
        self.combo_task_status.addItem('Архив', 'archived')
        self.combo_task_status.addItem('Все', 'all')
        self.combo_task_status.currentTextChanged.connect(self.load_tasks)
        status_layout.addWidget(self.combo_task_status)
        status_layout.addStretch()
        left_layout.addLayout(status_layout)
        task_sort_layout = QHBoxLayout()
        task_sort_layout.addWidget(QLabel('Сортировка:'))
        self.combo_task_sort = QComboBox()
        self.combo_task_sort.addItem('По дедлайну', 'due_date')
        self.combo_task_sort.addItem('По предмету', 'subject')
        self.combo_task_sort.addItem('По названию', 'title')
        self.combo_task_sort.currentTextChanged.connect(self.load_tasks)
        task_sort_layout.addWidget(self.combo_task_sort)
        self.btn_task_reverse = QPushButton('⇅')
        self.btn_task_reverse.setFixedWidth(30)
        self.btn_task_reverse.setToolTip('Изменить порядок')
        self.btn_task_reverse.clicked.connect(self.toggle_task_sort)
        task_sort_layout.addWidget(self.btn_task_reverse)
        task_sort_layout.addStretch()
        left_layout.addLayout(task_sort_layout)
        self.tasks_list = QListWidget()
        self.tasks_list.itemClicked.connect(self.show_task_details)
        left_layout.addWidget(self.tasks_list)
        btn_layout = QHBoxLayout()
        self.btn_add_task = QPushButton('Добавить')
        self.btn_edit_task = QPushButton('Ред.')
        self.btn_del_task = QPushButton('Удалить')
        self.btn_complete_task = QPushButton('Выполнено')
        self.btn_archive_task = QPushButton('В архив')
        self.btn_add_task.clicked.connect(self.add_task)
        self.btn_edit_task.clicked.connect(self.edit_task)
        self.btn_del_task.clicked.connect(self.delete_task)
        self.btn_complete_task.clicked.connect(self.toggle_complete_task)
        self.btn_archive_task.clicked.connect(self.toggle_archive_task)
        btn_layout.addWidget(self.btn_add_task)
        btn_layout.addWidget(self.btn_edit_task)
        btn_layout.addWidget(self.btn_del_task)
        btn_layout.addWidget(self.btn_complete_task)
        btn_layout.addWidget(self.btn_archive_task)
        left_layout.addLayout(btn_layout)
        right_layout = QVBoxLayout()
        right_layout.addWidget(QLabel('<h2>Детали задания</h2>'))
        info_frame = QFrame()
        info_frame.setFrameShape(QFrame.Shape.StyledPanel)
        info_layout = QVBoxLayout(info_frame)
        self.task_title = QLabel('Название: -')
        self.task_subject = QLabel('Предмет: -')
        self.task_date = QLabel('Срок: -')
        self.task_completed = QLabel('Выполнение: -')
        self.task_status = QLabel('Архив: -')
        info_layout.addWidget(self.task_title)
        info_layout.addWidget(self.task_subject)
        info_layout.addWidget(self.task_date)
        info_layout.addWidget(self.task_completed)
        info_layout.addWidget(self.task_status)
        info_layout.addSpacing(10)
        info_layout.addWidget(QLabel('<b>Описание:</b>'))
        self.task_desc = QLabel('-')
        self.task_desc.setWordWrap(True)
        self.task_desc.setAlignment(Qt.AlignmentFlag.AlignTop)
        info_layout.addWidget(self.task_desc)
        right_layout.addWidget(info_frame)
        right_layout.addStretch()
        layout.addLayout(left_layout, 30)
        layout.addLayout(right_layout, 70)

    def toggle_task_sort(self):
        logger.debug('toggle_task_sort')
        self.is_task_reverse = not self.is_task_reverse
        self.load_tasks()

    def load_tasks(self):
        logger.debug('load_tasks')
        try:
            self.tasks_list.clear()
            sort_mode = self.combo_task_sort.currentData()
            if not sort_mode:
                sort_mode = 'due_date'
            status_filter = self.combo_task_status.currentData()
            if not status_filter:
                status_filter = 'active'
            data = database.get_all_tasks(sort_by=sort_mode, reverse=self.is_task_reverse, status_filter=status_filter)
            for row in data:
                subject_name = row[3] if row[3] else 'Без предмета'
                date_str = row[2] if row[2] else 'Без даты'
                completed = row[5] if len(row) > 5 else 0
                status = row[4] if len(row) > 4 else 'active'
                status_text, color = self.get_task_status_info(row[2], completed, status)
                item_text = f'[{status_text}] [{subject_name}] {row[1]} ({date_str})'
                item = QListWidgetItem(item_text)
                item.setData(Qt.ItemDataRole.UserRole, row[0])
                item.setForeground(color)
                self.tasks_list.addItem(item)
        except Exception as e:
            logger.error(f'load_tasks: {e}', exc_info=True)

    def show_task_details(self, item):
        logger.debug(f'show_task_details: {item}')
        if not item:
            return
        try:
            tid = item.data(Qt.ItemDataRole.UserRole)
            if not tid:
                return
            data = database.get_task_details(tid)
            if not data:
                return
            self.task_title.setText(f'<b>Название:</b> {data[1]}')
            self.task_date.setText(f'<b>Срок:</b> {data[3] if data[3] else "Не указан"}')
            sub_id = data[4] if len(data) > 4 else None
            if sub_id:
                sub_details = database.get_subject_details(sub_id)
                sub_name = sub_details[1] if sub_details else 'Неизвестен'
                self.task_subject.setText(f'<b>Предмет:</b> {sub_name}')
            else:
                self.task_subject.setText('<b>Предмет:</b> -')
            status = data[5] if len(data) > 5 else 'active'
            completed = data[6] if len(data) > 6 else 0
            self.task_completed.setText(f'<b>Выполнение:</b> {"Да" if completed else "Нет"}')
            self.task_status.setText(f'<b>Архив:</b> {"Да" if status == "archived" else "Нет"}')
            if completed:
                self.btn_complete_task.setText('Не выполнено')
            else:
                self.btn_complete_task.setText('Выполнено')
            if status == 'archived':
                self.btn_archive_task.setText('Из архива')
            else:
                self.btn_archive_task.setText('В архив')
            self.task_desc.setText(data[2] if data[2] else 'Нет описания')
        except Exception as e:
            logger.error(f'show_task_details: {e}', exc_info=True)

    def add_task(self):
        logger.debug('add_task')
        try:
            subjects = database.get_all_subjects()
            dlg = TaskDialog(self, subjects)
            if dlg.exec():
                data = dlg.get_data()
                database.add_task(data['title'], data['description'], data['due_date'], data['subject_id'])
                self.load_tasks()
                self.update_calendar_deadlines()
        except Exception as e:
            logger.error(f'add_task: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', f'Произошла ошибка: {str(e)}')

    def edit_task(self):
        logger.debug('edit_task')
        item = self.tasks_list.currentItem()
        if not item:
            return
        try:
            tid = item.data(Qt.ItemDataRole.UserRole)
            data = database.get_task_details(tid)
            if not data:
                QMessageBox.warning(self, 'Ошибка', 'Задача не найдена')
                return
            subject_name = None
            sub_id = data[4] if len(data) > 4 else None
            if sub_id:
                sub = database.get_subject_details(sub_id)
                if sub:
                    subject_name = sub[1]
            task_data_for_dialog = (data[0], data[1], data[3], subject_name, data[2] if len(data) > 2 else '')
            subjects = database.get_all_subjects()
            dlg = TaskDialog(self, subjects, task_data_for_dialog)
            if dlg.exec():
                new_data = dlg.get_data()
                status = data[5] if len(data) > 5 else 'active'
                completed = data[6] if len(data) > 6 else 0
                database.update_task(tid, new_data['title'], new_data['description'], new_data['due_date'],
                                     new_data['subject_id'], status, completed)
                saved_tid = tid
                self.load_tasks()
                self.update_calendar_deadlines()
                for i in range(self.tasks_list.count()):
                    new_item = self.tasks_list.item(i)
                    if new_item.data(Qt.ItemDataRole.UserRole) == saved_tid:
                        self.tasks_list.setCurrentItem(new_item)
                        self.show_task_details(new_item)
                        break
        except Exception as e:
            logger.error(f'edit_task: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', f'Произошла ошибка: {str(e)}')

    def delete_task(self):
        logger.debug('delete_task')
        item = self.tasks_list.currentItem()
        if not item:
            return
        try:
            reply = QMessageBox.question(self, 'Удаление', 'Удалить задание?',
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.Yes:
                tid = item.data(Qt.ItemDataRole.UserRole)
                database.delete_task(tid)
                self.load_tasks()
                self.task_title.setText('Название: -')
                self.task_subject.setText('Предмет: -')
                self.task_date.setText('Срок: -')
                self.task_completed.setText('Выполнение: -')
                self.task_status.setText('Архив: -')
                self.task_desc.setText('-')
                self.update_calendar_deadlines()
        except Exception as e:
            logger.error(f'delete_task: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', f'Произошла ошибка: {str(e)}')

    def toggle_complete_task(self):
        logger.debug('toggle_complete_task')
        try:
            item = self.tasks_list.currentItem()
            if not item:
                return
            tid = item.data(Qt.ItemDataRole.UserRole)
            data = database.get_task_details(tid)
            if not data:
                return
            completed = data[6] if len(data) > 6 else 0
            new_completed = 0 if completed else 1
            database.toggle_task_completed(tid, new_completed)
            saved_tid = tid
            self.load_tasks()
            for i in range(self.tasks_list.count()):
                new_item = self.tasks_list.item(i)
                if new_item.data(Qt.ItemDataRole.UserRole) == saved_tid:
                    self.tasks_list.setCurrentItem(new_item)
                    self.show_task_details(new_item)
                    break
        except Exception as e:
            logger.error(f'toggle_complete_task: {e}', exc_info=True)

    def toggle_archive_task(self):
        logger.debug('toggle_archive_task')
        try:
            item = self.tasks_list.currentItem()
            if not item:
                return
            tid = item.data(Qt.ItemDataRole.UserRole)
            data = database.get_task_details(tid)
            if not data:
                return
            current_status = data[5] if len(data) > 5 else 'active'
            if current_status == 'archived':
                database.restore_task(tid)
                QMessageBox.information(self, 'Готово', 'Задание восстановлено из архива')
            else:
                database.archive_task(tid)
                QMessageBox.information(self, 'Готово', 'Задание перемещено в архив')
            saved_tid = tid
            self.load_tasks()
            self.update_calendar_deadlines()
            for i in range(self.tasks_list.count()):
                new_item = self.tasks_list.item(i)
                if new_item.data(Qt.ItemDataRole.UserRole) == saved_tid:
                    self.tasks_list.setCurrentItem(new_item)
                    self.show_task_details(new_item)
                    break
        except Exception as e:
            logger.error(f'toggle_archive_task: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', f'Произошла ошибка: {str(e)}')

    def setup_notes_ui(self):
        page = self.ui.page_notes
        layout = QVBoxLayout(page)
        layout.setContentsMargins(0, 0, 0, 0)
        self.lbl_current_file = QLabel('<b>Новый файл</b>')
        self.lbl_current_file.setStyleSheet('font-size: 12px; padding: 5px; color: gray;')
        layout.addWidget(self.lbl_current_file)
        toolbar_layout = QHBoxLayout()
        toolbar_layout.setContentsMargins(5, 5, 5, 5)
        self.btn_open = QPushButton('Открыть')
        self.btn_open.clicked.connect(self.open_file)
        toolbar_layout.addWidget(self.btn_open)
        self.btn_save = QPushButton('Сохранить')
        self.btn_save.clicked.connect(self.save_file)
        toolbar_layout.addWidget(self.btn_save)
        self.btn_save_as = QPushButton('Сохранить как...')
        self.btn_save_as.clicked.connect(self.save_file_as)
        toolbar_layout.addWidget(self.btn_save_as)
        toolbar_layout.addSpacing(15)
        toolbar_layout.addWidget(QLabel('Размер:'))
        self.combo_font_size = QComboBox()
        self.combo_font_size.addItems(['8', '9', '10', '11', '12', '14', '16', '18', '20', '24', '28', '36'])
        self.combo_font_size.setCurrentText('12')
        self.combo_font_size.currentTextChanged.connect(self.change_font_size)
        toolbar_layout.addWidget(self.combo_font_size)
        self.btn_align_center = QPushButton('Центр')
        self.btn_align_center.setCheckable(True)
        self.btn_align_center.clicked.connect(self.toggle_align_center)
        toolbar_layout.addWidget(self.btn_align_center)
        self.btn_bold = QPushButton('B')
        self.btn_bold.setCheckable(True)
        self.btn_bold.clicked.connect(self.toggle_bold)
        toolbar_layout.addWidget(self.btn_bold)
        self.btn_italic = QPushButton('I')
        self.btn_italic.setCheckable(True)
        self.btn_italic.clicked.connect(self.toggle_italic)
        toolbar_layout.addWidget(self.btn_italic)
        toolbar_layout.addStretch()
        layout.addLayout(toolbar_layout)
        self.notes_editor = QTextEdit()
        self.notes_editor.setPlaceholderText('Текст редактора...')
        layout.addWidget(self.notes_editor)

    def update_note_path_label(self):
        if self.current_note_path:
            self.lbl_current_file.setText(f'[Открыт] {self.current_note_path}')
        else:
            self.lbl_current_file.setText('<b>Новый файл</b> (не сохранен)')

    def change_font_size(self, size):
        fmt = self.notes_editor.currentCharFormat()
        fmt.setFontPointSize(float(size))
        self.notes_editor.setCurrentCharFormat(fmt)

    def toggle_align_center(self):
        if self.btn_align_center.isChecked():
            self.notes_editor.setAlignment(Qt.AlignmentFlag.AlignCenter)
        else:
            self.notes_editor.setAlignment(Qt.AlignmentFlag.AlignLeft)

    def toggle_bold(self):
        fmt = self.notes_editor.currentCharFormat()
        if self.btn_bold.isChecked():
            fmt.setFontWeight(QFont.Weight.Bold)
        else:
            fmt.setFontWeight(QFont.Weight.Normal)
        self.notes_editor.setCurrentCharFormat(fmt)

    def toggle_italic(self):
        fmt = self.notes_editor.currentCharFormat()
        fmt.setFontItalic(self.btn_italic.isChecked())
        self.notes_editor.setCurrentCharFormat(fmt)

    def open_file(self):
        logger.debug('open_file')
        try:
            start_dir = self.storage_path if self.storage_path else os.path.expanduser('~')
            filepath, _ = QFileDialog.getOpenFileName(self, 'Открыть файл', start_dir,
                                                      'Документы (*.docx *.txt);;Word (*.docx);;Текст (*.txt)')
            if filepath:
                self.load_file_content(filepath)
                self.current_note_path = filepath
                self.update_note_path_label()
        except Exception as e:
            logger.error(f'open_file: {e}', exc_info=True)

    def save_file(self):
        logger.debug('save_file')
        try:
            if self.current_note_path:
                self.save_file_content(self.current_note_path)
            else:
                self.save_file_as()
        except Exception as e:
            logger.error(f'save_file: {e}', exc_info=True)

    def save_file_as(self):
        logger.debug('save_file_as')
        try:
            start_dir = self.storage_path if self.storage_path else os.path.expanduser('~')
            filepath, _ = QFileDialog.getSaveFileName(self, 'Сохранить файл', start_dir,
                                                      'Word Documents (*.docx);;Text Files (*.txt)')
            if filepath:
                self.save_file_content(filepath)
                self.current_note_path = filepath
                self.update_note_path_label()
        except Exception as e:
            logger.error(f'save_file_as: {e}', exc_info=True)

    def load_file_content(self, path):
        logger.debug(f'load_file_content: {path}')
        try:
            if path.endswith('.docx'):
                doc = Document(path)
                self.notes_editor.clear()
                cursor = self.notes_editor.textCursor()
                is_first_paragraph = True
                for para in doc.paragraphs:
                    if not is_first_paragraph:
                        cursor.insertBlock()
                    else:
                        is_first_paragraph = False
                    block_fmt = QTextBlockFormat()
                    alignment = para.alignment
                    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                        block_fmt.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                        block_fmt.setAlignment(Qt.AlignmentFlag.AlignRight)
                    else:
                        block_fmt.setAlignment(Qt.AlignmentFlag.AlignLeft)
                    cursor.setBlockFormat(block_fmt)
                    for run in para.runs:
                        if run.text:
                            char_fmt = QTextCharFormat()
                            f_size = None
                            if run.font.size:
                                f_size = run.font.size.pt
                            if not f_size or f_size < 1:
                                f_size = 12
                            char_fmt.setFontPointSize(f_size)
                            if run.font.bold:
                                char_fmt.setFontWeight(QFont.Weight.Bold)
                            if run.font.italic:
                                char_fmt.setFontItalic(True)
                            cursor.insertText(run.text, char_fmt)
                    if not para.text.strip() and len(para.runs) == 0:
                        cursor.insertText(' ')
            elif path.endswith('.txt'):
                with open(path, 'r', encoding='utf-8') as f:
                    self.notes_editor.setPlainText(f.read())
            else:
                QMessageBox.warning(self, 'Ошибка', 'Неподдерживаемый формат')
        except Exception as e:
            logger.error(f'load_file_content: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', f'Не удалось прочитать файл: {e}')

    def save_file_content(self, path):
        logger.debug(f'save_file_content: {path}')
        try:
            if path.endswith('.docx'):
                doc = Document()
                block = self.notes_editor.document().begin()
                while block.isValid():
                    text = block.text()
                    char_fmt = block.charFormat()
                    block_fmt = block.blockFormat()
                    para = doc.add_paragraph(text, style=None)
                    if not para.runs:
                        run = para.add_run('')
                    else:
                        run = para.runs[0]
                    font_size = char_fmt.fontPointSize()
                    if font_size <= 1:
                        font_size = 11
                    run.font.size = Pt(font_size)
                    if char_fmt.fontWeight() == QFont.Weight.Bold:
                        run.font.bold = True
                    else:
                        run.font.bold = False
                    if char_fmt.fontItalic():
                        run.font.italic = True
                    else:
                        run.font.italic = False
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    if block_fmt.alignment() == Qt.AlignmentFlag.AlignCenter:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif block_fmt.alignment() == Qt.AlignmentFlag.AlignRight:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    block = block.next()
                doc.save(path)
                QMessageBox.information(self, 'Успех', f'Файл сохранен: {path}')
            elif path.endswith('.txt'):
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(self.notes_editor.toPlainText())
                QMessageBox.information(self, 'Успех', f'Файл сохранен: {path}')
            else:
                QMessageBox.warning(self, 'Ошибка', 'Неподдерживаемый формат')
        except Exception as e:
            logger.error(f'save_file_content: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', f'Не удалось сохранить: {e}')

    def setup_grades_ui(self):
        page = self.ui.page_grades
        layout = QVBoxLayout(page)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(15)

        top_layout = QHBoxLayout()
        top_layout.addWidget(QLabel('<h2>Журнал оценок</h2>'))
        top_layout.addStretch()
        top_layout.addWidget(QLabel('Семестр:'))
        self.combo_grade_semester = QComboBox()
        self.combo_grade_semester.addItem('Все', None)
        self.combo_grade_semester.currentIndexChanged.connect(self.load_grades)
        top_layout.addWidget(self.combo_grade_semester)
        layout.addLayout(top_layout)

        splitter = QSplitter(Qt.Orientation.Horizontal)

        left_widget = QWidget()
        left_layout = QVBoxLayout(left_widget)
        left_layout.setContentsMargins(0, 0, 0, 0)
        left_layout.addWidget(QLabel('<b>Список оценок</b>'))
        self.grades_table = QTableWidget()
        self.grades_table.setColumnCount(6)
        self.grades_table.setHorizontalHeaderLabels(['Предмет', 'Оценка', 'Тип', 'Дата', 'Семестр', 'Вес'])
        self.grades_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
        self.grades_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.grades_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.grades_table.itemClicked.connect(self.on_grade_selected)
        left_layout.addWidget(self.grades_table)
        btn_layout = QHBoxLayout()
        self.btn_add_grade = QPushButton('Добавить')
        self.btn_edit_grade = QPushButton('Изменить')
        self.btn_del_grade = QPushButton('Удалить')
        self.btn_add_grade.clicked.connect(self.add_grade)
        self.btn_edit_grade.clicked.connect(self.edit_grade)
        self.btn_del_grade.clicked.connect(self.delete_grade)
        btn_layout.addWidget(self.btn_add_grade)
        btn_layout.addWidget(self.btn_edit_grade)
        btn_layout.addWidget(self.btn_del_grade)
        left_layout.addLayout(btn_layout)
        splitter.addWidget(left_widget)

        right_widget = QWidget()
        right_layout = QVBoxLayout(right_widget)
        right_layout.setContentsMargins(0, 0, 0, 0)
        right_layout.addWidget(QLabel('<b>Статистика</b>'))

        self.stats_table = QTableWidget()
        self.stats_table.setColumnCount(6)
        self.stats_table.setHorizontalHeaderLabels(['Предмет', 'Среднее', 'Кол-во', 'Мин', 'Макс', 'Взвешенное'])
        # ИСПРАВЛЕНО: ресайз по контенту + растягивание последней колонки
        self.stats_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.ResizeToContents)
        self.stats_table.horizontalHeader().setStretchLastSection(True)
        self.stats_table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        right_layout.addWidget(self.stats_table)

        self.lbl_overall = QLabel('Общий средний балл: -')
        self.lbl_overall.setStyleSheet('font-size: 14px; padding: 10px;')
        right_layout.addWidget(self.lbl_overall)

        self.lbl_performance = QLabel('Успеваемость: -')
        self.lbl_performance.setStyleSheet('font-size: 14px; padding: 10px;')
        right_layout.addWidget(self.lbl_performance)

        right_layout.addStretch()
        splitter.addWidget(right_widget)
        splitter.setSizes([600, 400])
        layout.addWidget(splitter)

    def load_grades(self):
        logger.debug('load_grades')
        try:
            semester = self.combo_grade_semester.currentData()
            data = database.get_all_grades(semester=semester)
            self.grades_table.setRowCount(len(data))
            self.grades_map = {}
            for row_idx, row in enumerate(data):
                grade_id, subject_id, grade, grade_type, date, sem, weight, description, sub_name = row
                self.grades_table.setItem(row_idx, 0, QTableWidgetItem(sub_name or ''))
                grade_item = QTableWidgetItem(str(grade))
                grade_val = float(grade)
                if grade_val >= 4.5:
                    grade_item.setForeground(QBrush(QColor(0, 128, 0)))
                elif grade_val >= 3.5:
                    grade_item.setForeground(QBrush(QColor(0, 100, 200)))
                elif grade_val >= 2.5:
                    grade_item.setForeground(QBrush(QColor(200, 150, 0)))
                else:
                    grade_item.setForeground(QBrush(QColor(200, 0, 0)))
                self.grades_table.setItem(row_idx, 1, grade_item)
                self.grades_table.setItem(row_idx, 2, QTableWidgetItem(grade_type or ''))
                self.grades_table.setItem(row_idx, 3, QTableWidgetItem(date or ''))
                self.grades_table.setItem(row_idx, 4, QTableWidgetItem(sem or ''))
                self.grades_table.setItem(row_idx, 5, QTableWidgetItem(str(weight or 1.0)))
                self.grades_map[row_idx] = grade_id
            self.load_statistics()
            self.update_semester_combo()
        except Exception as e:
            logger.error(f'load_grades: {e}', exc_info=True)

    def update_semester_combo(self):
        current = self.combo_grade_semester.currentData()
        self.combo_grade_semester.blockSignals(True)
        self.combo_grade_semester.clear()
        self.combo_grade_semester.addItem('Все', None)
        semesters = database.get_semesters()
        for sem in semesters:
            self.combo_grade_semester.addItem(sem, sem)
        if current:
            idx = self.combo_grade_semester.findData(current)
            if idx >= 0:
                self.combo_grade_semester.setCurrentIndex(idx)
        self.combo_grade_semester.blockSignals(False)

    def load_statistics(self):
        logger.debug('load_statistics')
        try:
            semester = self.combo_grade_semester.currentData()
            stats = GradeStatistics()
            data = stats.load_statistics(semester=semester)
            self.stats_table.setRowCount(len(data))
            for row_idx, row in enumerate(data):
                sub_name, avg_grade, count, min_grade, max_grade, weighted_avg = row
                self.stats_table.setItem(row_idx, 0, QTableWidgetItem(sub_name or ''))
                self.stats_table.setItem(row_idx, 1, QTableWidgetItem(f'{avg_grade:.2f}' if avg_grade else '-'))
                self.stats_table.setItem(row_idx, 2, QTableWidgetItem(str(count or 0)))
                self.stats_table.setItem(row_idx, 3, QTableWidgetItem(str(min_grade) if min_grade else '-'))
                self.stats_table.setItem(row_idx, 4, QTableWidgetItem(str(max_grade) if max_grade else '-'))
                self.stats_table.setItem(row_idx, 5, QTableWidgetItem(f'{weighted_avg:.2f}' if weighted_avg else '-'))
            overall = stats.get_overall_average(semester=semester)
            self.lbl_overall.setText(f'Общий средний балл: <b>{overall:.2f}</b>')
            level, color = stats.get_performance_level(overall)
            self.lbl_performance.setText(f'Успеваемость: <b>{level}</b>')
            self.lbl_performance.setStyleSheet(f'font-size: 14px; padding: 10px; color: {color.name()};')
        except Exception as e:
            logger.error(f'load_statistics: {e}', exc_info=True)

    def on_grade_selected(self, item):
        row = item.row()
        if row in self.grades_map:
            self.selected_grade_id = self.grades_map[row]

    def add_grade(self):
        logger.debug('add_grade')
        try:
            subjects = database.get_all_subjects()
            if not subjects:
                QMessageBox.warning(self, 'Ошибка', 'Сначала добавьте предметы')
                return
            dlg = GradeDialog(self, subjects)
            if dlg.exec():
                data = dlg.get_data()
                database.add_grade(**data)
                self.load_grades()
        except Exception as e:
            logger.error(f'add_grade: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', f'Произошла ошибка: {str(e)}')

    def edit_grade(self):
        logger.debug('edit_grade')
        try:
            selected = self.grades_table.selectedItems()
            if not selected:
                QMessageBox.warning(self, 'Ошибка', 'Выберите оценку')
                return
            row = selected[0].row()
            if row not in self.grades_map:
                return
            grade_id = self.grades_map[row]
            data = database.get_grade_by_id(grade_id)
            if not data:
                return
            subjects = database.get_all_subjects()
            dlg = GradeDialog(self, subjects, data)
            if dlg.exec():
                new_data = dlg.get_data()
                database.update_grade(grade_id, **new_data)
                self.load_grades()
        except Exception as e:
            logger.error(f'edit_grade: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', f'Произошла ошибка: {str(e)}')

    def delete_grade(self):
        logger.debug('delete_grade')
        try:
            selected = self.grades_table.selectedItems()
            if not selected:
                QMessageBox.warning(self, 'Ошибка', 'Выберите оценку')
                return
            row = selected[0].row()
            if row not in self.grades_map:
                return
            grade_id = self.grades_map[row]
            reply = QMessageBox.question(self, 'Удаление', 'Удалить оценку?',
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.Yes:
                database.delete_grade(grade_id)
                self.load_grades()
        except Exception as e:
            logger.error(f'delete_grade: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', f'Произошла ошибка: {str(e)}')

    def setup_wishlist_ui(self):
        page = self.ui.page_wishlist
        layout = QHBoxLayout(page)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(20)

        left_layout = QVBoxLayout()
        left_layout.addWidget(QLabel('<h2>Список желаний</h2>'))

        filter_layout = QHBoxLayout()
        filter_layout.addWidget(QLabel('Статус:'))
        self.combo_wish_status = QComboBox()
        self.combo_wish_status.addItem('Все', 'all')
        self.combo_wish_status.addItem('Активные', 'pending')
        self.combo_wish_status.addItem('Выполненные', 'done')
        self.combo_wish_status.currentTextChanged.connect(self.load_wishlist)
        filter_layout.addWidget(self.combo_wish_status)

        filter_layout.addWidget(QLabel('Категория:'))
        self.combo_wish_category = QComboBox()
        self.combo_wish_category.addItem('Все', 'all')
        self.combo_wish_category.currentTextChanged.connect(self.load_wishlist)
        filter_layout.addWidget(self.combo_wish_category)
        filter_layout.addStretch()
        left_layout.addLayout(filter_layout)

        sort_layout = QHBoxLayout()
        sort_layout.addWidget(QLabel('Сортировка:'))
        self.combo_wish_sort = QComboBox()
        self.combo_wish_sort.addItem('По дате', 'created_at')
        self.combo_wish_sort.addItem('По названию', 'title')
        self.combo_wish_sort.addItem('По приоритету', 'priority')
        self.combo_wish_sort.currentTextChanged.connect(self.load_wishlist)
        sort_layout.addWidget(self.combo_wish_sort)
        sort_layout.addStretch()
        left_layout.addLayout(sort_layout)

        self.wishlist_list = QListWidget()
        self.wishlist_list.itemClicked.connect(self.show_wishlist_detail)
        left_layout.addWidget(self.wishlist_list)

        btn_layout = QHBoxLayout()
        self.btn_add_wish = QPushButton('Добавить')
        self.btn_edit_wish = QPushButton('Изменить')
        self.btn_del_wish = QPushButton('Удалить')
        self.btn_done_wish = QPushButton('Выполнено')
        self.btn_add_wish.clicked.connect(self.add_wishlist_item)
        self.btn_edit_wish.clicked.connect(self.edit_wishlist_item)
        self.btn_del_wish.clicked.connect(self.delete_wishlist_item)
        self.btn_done_wish.clicked.connect(self.toggle_wishlist_done)
        btn_layout.addWidget(self.btn_add_wish)
        btn_layout.addWidget(self.btn_edit_wish)
        btn_layout.addWidget(self.btn_del_wish)
        btn_layout.addWidget(self.btn_done_wish)
        left_layout.addLayout(btn_layout)

        right_layout = QVBoxLayout()
        right_layout.addWidget(QLabel('<h2>Детали</h2>'))
        info_frame = QFrame()
        info_frame.setFrameShape(QFrame.Shape.StyledPanel)
        info_layout = QVBoxLayout(info_frame)
        self.lbl_wish_title = QLabel('Название: -')
        self.lbl_wish_desc = QLabel('Описание: -')
        self.lbl_wish_category = QLabel('Категория: -')
        self.lbl_wish_priority = QLabel('Приоритет: -')
        self.lbl_wish_status = QLabel('Статус: -')
        self.lbl_wish_date = QLabel('Дата: -')
        info_layout.addWidget(self.lbl_wish_title)
        info_layout.addWidget(self.lbl_wish_desc)
        info_layout.addWidget(self.lbl_wish_category)
        info_layout.addWidget(self.lbl_wish_priority)
        info_layout.addWidget(self.lbl_wish_status)
        info_layout.addWidget(self.lbl_wish_date)
        right_layout.addWidget(info_frame)
        right_layout.addStretch()

        layout.addLayout(left_layout, 35)
        layout.addLayout(right_layout, 65)

    def load_wishlist(self):
        logger.debug('load_wishlist')
        try:
            self.wishlist_list.clear()
            status_filter = self.combo_wish_status.currentData()
            category_filter = self.combo_wish_category.currentData()
            sort_by = self.combo_wish_sort.currentData()

            data = database.get_all_wishlist(
                sort_by=sort_by,
                status_filter=status_filter,
                category_filter=category_filter
            )

            categories = set()
            for row in data:
                item_id, title, description, category, priority, status, created_at = row
                if category:
                    categories.add(category)
                priority_display = {'high': 'HIGH', 'medium': 'MED', 'low': 'LOW'}.get(priority, priority)
                status_mark = 'V' if status == 'done' else ' '
                display = f'[{status_mark}] [{priority_display}] {title}'
                item = QListWidgetItem(display)
                item.setData(Qt.ItemDataRole.UserRole, item_id)
                if status == 'done':
                    item.setForeground(QColor(128, 128, 128))
                elif priority == 'high':
                    item.setForeground(QColor(220, 20, 60))
                elif priority == 'low':
                    item.setForeground(QColor(0, 128, 0))
                self.wishlist_list.addItem(item)

            current_cat = self.combo_wish_category.currentData()
            self.combo_wish_category.blockSignals(True)
            self.combo_wish_category.clear()
            self.combo_wish_category.addItem('Все', 'all')
            for cat in sorted(categories):
                self.combo_wish_category.addItem(cat, cat)
            if current_cat:
                idx = self.combo_wish_category.findData(current_cat)
                if idx >= 0:
                    self.combo_wish_category.setCurrentIndex(idx)
            self.combo_wish_category.blockSignals(False)

        except Exception as e:
            logger.error(f'load_wishlist: {e}', exc_info=True)

    def show_wishlist_detail(self, item):
        if not item:
            return
        try:
            item_id = item.data(Qt.ItemDataRole.UserRole)
            data = database.get_wishlist_by_id(item_id)
            if not data:
                return
            item_id, title, description, category, priority, status, created_at = data
            self.lbl_wish_title.setText(f'<b>Название:</b> {title or "-"}')
            self.lbl_wish_desc.setText(f'<b>Описание:</b> {description or "-"}')
            self.lbl_wish_desc.setWordWrap(True)
            cat_display = category or 'Другое'
            self.lbl_wish_category.setText(f'<b>Категория:</b> {cat_display}')
            priority_display = {'high': 'Высокий', 'medium': 'Средний', 'low': 'Низкий'}.get(priority, priority)
            self.lbl_wish_priority.setText(f'<b>Приоритет:</b> {priority_display}')
            self.lbl_wish_status.setText(f'<b>Статус:</b> {"Выполнено" if status == "done" else "Активно"}')
            self.lbl_wish_date.setText(f'<b>Дата:</b> {created_at or "-"}')
            if status == 'done':
                self.btn_done_wish.setText('Восстановить')
            else:
                self.btn_done_wish.setText('Выполнено')
        except Exception as e:
            logger.error(f'show_wishlist_detail: {e}', exc_info=True)

    def add_wishlist_item(self):
        try:
            dlg = WishlistDialog(self)
            if dlg.exec():
                data = dlg.get_data()
                if not data['title']:
                    QMessageBox.warning(self, 'Ошибка', 'Название не может быть пустым')
                    return
                database.add_wishlist_item(data['title'], data['description'],
                                          data['category'], data['priority'])
                self.load_wishlist()
        except Exception as e:
            logger.error(f'add_wishlist_item: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', str(e))

    def edit_wishlist_item(self):
        item = self.wishlist_list.currentItem()
        if not item:
            return
        try:
            item_id = item.data(Qt.ItemDataRole.UserRole)
            data = database.get_wishlist_by_id(item_id)
            if not data:
                return
            dlg = WishlistDialog(self, data)
            if dlg.exec():
                new_data = dlg.get_data()
                database.update_wishlist_item(item_id, new_data['title'], new_data['description'],
                                             new_data['category'], new_data['priority'], data[5])
                self.load_wishlist()
        except Exception as e:
            logger.error(f'edit_wishlist_item: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', str(e))

    def delete_wishlist_item(self):
        item = self.wishlist_list.currentItem()
        if not item:
            return
        try:
            item_id = item.data(Qt.ItemDataRole.UserRole)
            reply = QMessageBox.question(self, 'Удаление', 'Удалить желание?',
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.Yes:
                database.delete_wishlist_item(item_id)
                self.load_wishlist()
                self.lbl_wish_title.setText('Название: -')
                self.lbl_wish_desc.setText('Описание: -')
        except Exception as e:
            logger.error(f'delete_wishlist_item: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', str(e))

    def toggle_wishlist_done(self):
        item = self.wishlist_list.currentItem()
        if not item:
            return
        try:
            item_id = item.data(Qt.ItemDataRole.UserRole)
            data = database.get_wishlist_by_id(item_id)
            if not data:
                return
            new_status = 'pending' if data[5] == 'done' else 'done'
            database.update_wishlist_item(item_id, data[1], data[2], data[3], data[4], new_status)
            self.load_wishlist()
            for i in range(self.wishlist_list.count()):
                new_item = self.wishlist_list.item(i)
                if new_item.data(Qt.ItemDataRole.UserRole) == item_id:
                    self.wishlist_list.setCurrentItem(new_item)
                    self.show_wishlist_detail(new_item)
                    break
        except Exception as e:
            logger.error(f'toggle_wishlist_done: {e}', exc_info=True)
            QMessageBox.critical(self, 'Ошибка', str(e))

    def setup_notifications_menu(self):
        notif_menu = self.ui.menubar.addMenu('Уведомления')
        check_action = notif_menu.addAction('Проверить дедлайны')
        check_action.triggered.connect(self.notification_manager.manual_check)
        notif_menu.addSeparator()
        settings_action = notif_menu.addAction('Настройки уведомлений')
        settings_action.triggered.connect(self.show_notification_settings)

    def show_notification_settings(self):
        QMessageBox.information(self, 'Настройки уведомлений',
                                'Уведомления проверяются автоматически каждые 30 минут. Срочные уведомления показываются в трее. Используйте меню для ручной проверки.')

    def closeEvent(self, event):
        if self.notification_manager and self.notification_manager.tray_icon:
            self.hide()
            self.notification_manager.show_tray_message('Study Planner',
                                                        'Приложение свернуто в трей. Двойной клик для разворачивания.')
            event.ignore()
        else:
            if self.notification_manager:
                self.notification_manager.cleanup()
            event.accept()


if __name__ == '__main__':
    app = QApplication(sys.argv)
    window = MainWindow()
    window.show()
    sys.exit(app.exec())
