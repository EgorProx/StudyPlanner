import logging
import os
from PyQt6.QtWidgets import (QVBoxLayout, QHBoxLayout, QLabel, QPushButton,
                             QComboBox, QTextEdit, QFileDialog, QMessageBox)
from PyQt6.QtCore import Qt
from PyQt6.QtGui import QFont
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)


class NotesPage:
    def __init__(self, mw):
        self.mw = mw
        self._setup_ui()

    def _setup_ui(self):
        page = self.mw.ui.page_notes
        layout = QVBoxLayout(page)
        layout.setContentsMargins(0, 0, 0, 0)

        self.lbl_current_file = QLabel('<b>Новый файл</b>')
        self.lbl_current_file.setStyleSheet('font-size: 12px; padding: 5px; color: gray;')
        layout.addWidget(self.lbl_current_file)

        toolbar_layout = QHBoxLayout()
        toolbar_layout.setContentsMargins(5, 5, 5, 5)
        self.btn_open = QPushButton('Открыть')
        self.btn_open.clicked.connect(self._open)
        toolbar_layout.addWidget(self.btn_open)
        self.btn_save = QPushButton('Сохранить')
        self.btn_save.clicked.connect(self._save)
        toolbar_layout.addWidget(self.btn_save)
        self.btn_save_as = QPushButton('Сохранить как...')
        self.btn_save_as.clicked.connect(self._save_as)
        toolbar_layout.addWidget(self.btn_save_as)
        toolbar_layout.addSpacing(15)
        toolbar_layout.addWidget(QLabel('Размер:'))
        self.combo_font_size = QComboBox()
        self.combo_font_size.addItems(['8', '9', '10', '11', '12', '14', '16', '18', '20', '24', '28', '36'])
        self.combo_font_size.setCurrentText('12')
        self.combo_font_size.currentTextChanged.connect(self._change_font_size)
        toolbar_layout.addWidget(self.combo_font_size)
        self.btn_align_center = QPushButton('Центр')
        self.btn_align_center.setCheckable(True)
        self.btn_align_center.clicked.connect(self._toggle_align_center)
        toolbar_layout.addWidget(self.btn_align_center)
        self.btn_bold = QPushButton('B')
        self.btn_bold.setCheckable(True)
        self.btn_bold.clicked.connect(self._toggle_bold)
        toolbar_layout.addWidget(self.btn_bold)
        self.btn_italic = QPushButton('I')
        self.btn_italic.setCheckable(True)
        self.btn_italic.clicked.connect(self._toggle_italic)
        toolbar_layout.addWidget(self.btn_italic)
        toolbar_layout.addStretch()
        layout.addLayout(toolbar_layout)

        self.editor = QTextEdit()
        self.editor.setPlaceholderText('Текст редактора...')
        layout.addWidget(self.editor)

    def update_path_label(self):
        if self.mw.current_note_path:
            self.lbl_current_file.setText(f'[Открыт] {self.mw.current_note_path}')
        else:
            self.lbl_current_file.setText('<b>Новый файл</b> (не сохранен)')

    def _change_font_size(self, size):
        fmt = self.editor.currentCharFormat()
        fmt.setFontPointSize(float(size))
        self.editor.setCurrentCharFormat(fmt)

    def _toggle_align_center(self):
        if self.btn_align_center.isChecked():
            self.editor.setAlignment(Qt.AlignmentFlag.AlignCenter)
        else:
            self.editor.setAlignment(Qt.AlignmentFlag.AlignLeft)

    def _toggle_bold(self):
        fmt = self.editor.currentCharFormat()
        if self.btn_bold.isChecked():
            fmt.setFontWeight(QFont.Weight.Bold)
        else:
            fmt.setFontWeight(QFont.Weight.Normal)
        self.editor.setCurrentCharFormat(fmt)

    def _toggle_italic(self):
        fmt = self.editor.currentCharFormat()
        fmt.setFontItalic(self.btn_italic.isChecked())
        self.editor.setCurrentCharFormat(fmt)

    def _open(self):
        try:
            start_dir = self.mw.storage_path if self.mw.storage_path else os.path.expanduser('~')
            filepath, _ = QFileDialog.getOpenFileName(self.mw, 'Открыть файл', start_dir,
                                                       'Документы (*.docx *.txt);;Word (*.docx);;Текст (*.txt)')
            if filepath:
                self._load_file_content(filepath)
                self.mw.current_note_path = filepath
                self.update_path_label()
        except Exception as e:
            logger.error(f'open_file: {e}', exc_info=True)

    def _save(self):
        try:
            if self.mw.current_note_path:
                self._save_file_content(self.mw.current_note_path)
            else:
                self._save_as()
        except Exception as e:
            logger.error(f'save_file: {e}', exc_info=True)

    def _save_as(self):
        try:
            start_dir = self.mw.storage_path if self.mw.storage_path else os.path.expanduser('~')
            filepath, _ = QFileDialog.getSaveFileName(self.mw, 'Сохранить файл', start_dir,
                                                       'Word Documents (*.docx);;Text Files (*.txt)')
            if filepath:
                self._save_file_content(filepath)
                self.mw.current_note_path = filepath
                self.update_path_label()
        except Exception as e:
            logger.error(f'save_file_as: {e}', exc_info=True)

    def _load_file_content(self, path):
        try:
            if path.endswith('.docx'):
                doc = Document(path)
                self.editor.clear()
                cursor = self.editor.textCursor()
                is_first = True
                for para in doc.paragraphs:
                    if not is_first:
                        cursor.insertBlock()
                    else:
                        is_first = False
                    from PyQt6.QtGui import QTextBlockFormat, QTextCharFormat
                    block_fmt = QTextBlockFormat()
                    alignment = para.alignment
                    if alignment == WD_ALIGN_PARAGRAPH.CENTER:
                        block_fmt.setAlignment(Qt.AlignmentFlag.AlignCenter)
                    elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                        block_fmt.setAlignment(Qt.AlignmentFlag.AlignRight)
                    else:
                        block_fmt.setAlignment(Qt.AlignmentFlag.AlignLeft)
                    cursor.setBlockFormat(block_fmt)
                    for run in para.runs:
                        if run.text:
                            char_fmt = QTextCharFormat()
                            f_size = run.font.size.pt if run.font.size else 12
                            if not f_size or f_size < 1:
                                f_size = 12
                            char_fmt.setFontPointSize(f_size)
                            if run.font.bold:
                                char_fmt.setFontWeight(QFont.Weight.Bold)
                            if run.font.italic:
                                char_fmt.setFontItalic(True)
                            cursor.insertText(run.text, char_fmt)
                    if not para.text.strip() and len(para.runs) == 0:
                        cursor.insertText(' ')
            elif path.endswith('.txt'):
                with open(path, 'r', encoding='utf-8') as f:
                    self.editor.setPlainText(f.read())
            else:
                QMessageBox.warning(self.mw, 'Ошибка', 'Неподдерживаемый формат')
        except Exception as e:
            logger.error(f'load_file_content: {e}', exc_info=True)
            QMessageBox.critical(self.mw, 'Ошибка', f'Не удалось прочитать файл: {e}')

    def _save_file_content(self, path):
        try:
            if path.endswith('.docx'):
                doc = Document()
                block = self.editor.document().begin()
                while block.isValid():
                    text = block.text()
                    char_fmt = block.charFormat()
                    block_fmt = block.blockFormat()
                    para = doc.add_paragraph(text, style=None)
                    if not para.runs:
                        run = para.add_run('')
                    else:
                        run = para.runs[0]
                    font_size = char_fmt.fontPointSize()
                    if font_size <= 1:
                        font_size = 11
                    run.font.size = Pt(font_size)
                    run.font.bold = char_fmt.fontWeight() == QFont.Weight.Bold
                    run.font.italic = char_fmt.fontItalic()
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    if block_fmt.alignment() == Qt.AlignmentFlag.AlignCenter:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif block_fmt.alignment() == Qt.AlignmentFlag.AlignRight:
                        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    block = block.next()
                doc.save(path)
                QMessageBox.information(self.mw, 'Успех', f'Файл сохранен: {path}')
            elif path.endswith('.txt'):
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(self.editor.toPlainText())
                QMessageBox.information(self.mw, 'Успех', f'Файл сохранен: {path}')
            else:
                QMessageBox.warning(self.mw, 'Ошибка', 'Неподдерживаемый формат')
        except Exception as e:
            logger.error(f'save_file_content: {e}', exc_info=True)
            QMessageBox.critical(self.mw, 'Ошибка', f'Не удалось сохранить: {e}')
